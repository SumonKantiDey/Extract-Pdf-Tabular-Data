# Created by Avinashkumar


# lib for working with APIq
import convertapi
import time
import glob, os

# lib for working with the docx files
from docx import Document
import numpy as np
import pandas as pd
import io
import csv
from docx import Document

# -----------------------------------------
secret_key = 'oGlQf23cDqDTv6FO'

Tokens='9Irk2xSo'
convertapi.api_secret = secret_key

# -----------------------------------------

# ============
#  Functions
# ============
def con2docx(input_path):
    start = time.time()
    print("=> initialized docx convertion")
    print("start : ", start)
    result = convertapi.convert('docx', {'File': input_path})
    result.file.save(input_path.replace('.pdf', '.docx'))
    done = time.time()
    elapsed = done - start
    print("=> {0}/{1} || ".format(1, 1) + input_path + "  : " + str(elapsed))


# end of con2codx

def read_docx_tables(filename, tab_id=None, **kwargs):
    """
    parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

    Parameters:
        filename:   file name of a Word Document

        tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                    When [None] - return a list of DataFrames (parse all tables)

        kwargs:     arguments to pass to `pd.read_csv()` function

    Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
    """

    def read_docx_tab(tab, **kwargs):
        vf = io.StringIO()
        writer = csv.writer(vf)
        for row in tab.rows:
            writer.writerow(cell.text for cell in row.cells)
        vf.seek(0)
        return pd.read_csv(vf, **kwargs)

    doc = Document(filename)
    if tab_id is None:
        return [read_docx_tab(tab, **kwargs) for tab in doc.tables]
    else:
        try:
            return read_docx_tab(doc.tables[tab_id], **kwargs)
        except IndexError:
            print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
            raise
    # end of read_docx_tables


def prob_1(dfs):
    # here the last two tables are belongs to STOCK ASSESSED TABLE
    pre_check = dfs[len(dfs) - 1]
    pre_check.columns.tolist()

    content_in_the_headers = ""
    for i in pre_check.columns.tolist():
        content_in_the_headers += i

    # skip the code if the condition not satisified
    if 'Assessed' not in content_in_the_headers:
        print("prob 1 occur")
        return True
    return False


def prob_2(dfs):
    no_of_cln = len(dfs[len(dfs) - 2].columns.tolist())
    # skip the code if the condition not satisified
    if no_of_cln == 3:  # if the third last table column count is 3 the problem 2 occur
        print("prob 2 occur")
        return True
    return False


def TIME_ASSESSMENT(dfs):
    '''
        # creates 'TIME ASSESSED TABLE dataframe '

        returns the dataframe
    '''
    if prob_2(dfs):
        return None

    if not prob_1(dfs):
        # get the last table
        tb = dfs[len(dfs) - 1]
        tb = tb.drop(columns=tb.columns.tolist()[0])
        res = tb.loc[0, :]
        content = res.tolist()
        head = tb.columns.tolist()

        final = []
        for i, x in enumerate(head):
            final.append(x + " " + content[i])
        # final is variable in which the columns name will be stored
        # this operations are performed for renaming the dataframe column correctly

        copy = final
        for i, x in enumerate(copy):
            if "Unnamed:" in x:
                final[i] = copy[i][10:]

        tb.columns = final  # renaming the table columns
    else:
        tb = dfs[len(dfs) - 1]
        tb = tb.drop(['Unnamed: 0'], axis=1)
        hd = tb.iloc[0].tolist()
        hd[2] = ""
        hd[4] = ""

        tb.columns = hd
        tb = tb.iloc[1:]

    tb = tb.dropna()  # removing the empty row

    # genrating the list
    # which is need to be send as one of the parameter to function chdf()
    list_data = []
    for i, x in enumerate(tb.index):
        list_data.append([])
        list_data[i] = tb.iloc[i].tolist()

    # chdf refers to CHange Data Frame
    def chdf(list, header, df):
        # Gets three parameter as input
        # they are :> df_converted_as_list , df_columns_as_list , data_frame

        df_copy = df
        df_li = list
        for i, row in enumerate(df_li):
            if row[0] in header[0]:
                print("deleting index ", i)
                df_copy = df_copy.drop([i])
        return df_copy  # returns the finally created dataframe

    # after performing some of the drop function in the df.
    # the index of the dataframe will be not in order
    # so reframing the index of the dataframe

    tb.index = range(len(tb.index))

    final_df = chdf(list_data, tb.columns.tolist(), tb)
    final_df.index = range(len(final_df.index))  # reframing the index of data frame
    return final_df


# end of TIME_ASSESSMENT

def STOCK_ASSESSED_WEEK(dfs):
    count = 0
    if prob_2(dfs):
        count = 1

    return dfs[len(dfs) - 3 + count]


# end of STOCK_ASSESSED_WEEK

def STOCK_ASSESSED_MONTH(dfs):
    count = 0
    if prob_2(dfs):
        count = 1

    return dfs[len(dfs) - 2 + count]


# end of STOCK_ASSESSED_MONTH

def TIME_FIXTURES(dfs):
    count = 0
    if prob_2(dfs):
        count = 1

    tables = dfs[1:len(dfs) - 3 + count]

    # obtain only the headers
    table_headers = []
    for i, table in enumerate(tables):
        table_headers.append([])
        table_headers[i] = table.columns.tolist()

    TIME_FIXTURES = None
    for i, table_header in enumerate(table_headers):
        if 'Period' in table_header:
            TIME_FIXTURES = dfs[i + 1]
            break

    return TIME_FIXTURES


# end of TIME_FIXTURES

def STOCK_FIXTURES(dfs):
    count = 0
    if prob_2(dfs):
        count = 1

    tables = dfs[1:len(dfs) - 3 + count]
    # let omitt the first and last table
    # because first table will be not consider and last one is always TIME_ASSESSMENT_TABLE
    # obtain only the headers
    table_headers = []
    for i, table in enumerate(tables):
        table_headers.append([])
        table_headers[i] = table.columns.tolist()

    noted_index = 0
    for i, table_header in enumerate(table_headers):
        if 'Period' in table_header:
            noted_index = i + 1
            break

    table_indexes = list(range(1, len(dfs) - 3 + count))
    if noted_index in table_indexes:
        table_indexes.remove(noted_index)

    copy_headers = table_headers.copy()

    table_headers = []
    for i, index in enumerate(table_indexes):
        table_headers.append([])
        table_headers[i] = copy_headers[i]

    # table_indexes
    # table_headers

    tables_to_be_merged = []
    for it, i in enumerate(table_indexes):
        for j in table_indexes[it + 1:len(table_indexes)]:
            if dfs[i].columns.tolist() == dfs[j].columns.tolist():
                tables_to_be_merged.append([i, j])
    tmp = tables_to_be_merged

    STOCK_FIXTURES = []
    NEW_DF_INDEX = []
    for x in tmp:
        frames = []
        for tb_i in x:
            frames.append(dfs[tb_i])
        new = pd.concat(frames)
        NEW_DF_INDEX.append(x[0])
        STOCK_FIXTURES.append(new)

    # len(STOCK_FIXTURES)

    new_index = []
    for x in tmp:
        new_index += x
    for x in table_indexes:
        if x not in new_index:
            STOCK_FIXTURES.append(dfs[x])

    return STOCK_FIXTURES


# end of STOCK_FIXTURES

# ---------------------------------------------------------------------------------------------------------

# convert-code

# main-code
# The table extraction code starts here
# you have to give the input parameter as path of the ' docx file '

# input to docx path


def MAIN_CODE(path):
    # r"C:\Users\avina\Desktop\blue scheme\final code\docx files\Weekly Market Report 010319.docx"
    con2docx(path)
    pth = "."+path.split(".")[1] + ".docx"
    # dfs global variable
    print(pth)
    dfs = read_docx_tables(pth)
    TIME_ASSESSMENT_TABLE=[]
    STOCK_ASSESSED_WEEK_TABLE=[]
    STOCK_ASSESSED_MONTH_TABLE=[]
    TIME_FIXTURES_TABLE=[]
    STOCK_FIXTURES_TABLE=[]
    try:
        TIME_ASSESSMENT_TABLE = TIME_ASSESSMENT(dfs)
        STOCK_ASSESSED_WEEK_TABLE = STOCK_ASSESSED_WEEK(dfs)
        STOCK_ASSESSED_MONTH_TABLE = STOCK_ASSESSED_MONTH(dfs)
        TIME_FIXTURES_TABLE = TIME_FIXTURES(dfs)
        STOCK_FIXTURES_TABLE = STOCK_FIXTURES(dfs)
    except Exception as e:
        print(e)

    # TOTALLY 5 tables
    return TIME_ASSESSMENT_TABLE, TIME_FIXTURES_TABLE, STOCK_ASSESSED_MONTH_TABLE, STOCK_ASSESSED_WEEK_TABLE, STOCK_FIXTURES_TABLE

#MAIN_CODE('/home/blueschemeai/project/textextraction/apis/files/WeeklyMarketReport310519.pdf')